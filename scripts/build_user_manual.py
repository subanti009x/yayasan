from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("C:/yayasansekolah")
ASSET_DIR = ROOT / "manual-assets"
OUT_DIR = ROOT / "deliverables"
DOCX_PATH = OUT_DIR / "Manual_Pengguna_Teknis_Yayasan_Cendekia.docx"

OUT_DIR.mkdir(exist_ok=True)


SCREENSHOTS = {
    "home": ("01-beranda-dashboard.png", "Gambar 1. Dashboard utama halaman Beranda."),
    "unit_cards": ("02-beranda-unit-sekolah.png", "Gambar 2. Daftar unit sekolah pada halaman Beranda."),
    "unit_hero": ("03-halaman-unit-hero.png", "Gambar 3. Halaman profil unit sekolah dan tombol pendaftaran."),
    "registration": ("04-form-pendaftaran.png", "Gambar 4. Formulir pendaftaran online pada halaman unit sekolah."),
    "articles": ("05-artikel.png", "Gambar 5. Halaman artikel sekolah."),
    "faq": ("06-faq.png", "Gambar 6. Halaman FAQ pendaftaran."),
    "contact": ("07-kontak.png", "Gambar 7. Halaman kontak dan peta lokasi."),
    "admin_login": ("08-admin-login.png", "Gambar 8. Form login dashboard admin."),
    "admin_articles": ("09-admin-artikel.png", "Gambar 9. Modul pengelolaan artikel di dashboard admin."),
    "admin_branding": ("10-admin-branding.png", "Gambar 10. Modul Branding & Tampilan."),
    "admin_identity": ("11-admin-identitas.png", "Gambar 11. Modul Identitas & Kontak Yayasan."),
    "admin_school_list": ("12-admin-unit-sekolah.png", "Gambar 12. Daftar unit sekolah di dashboard admin."),
    "admin_school_edit": ("13-admin-edit-unit.png", "Gambar 13. Form edit konten unit sekolah."),
    "admin_faq": ("14-admin-faq.png", "Gambar 14. Modul pengelolaan FAQ pendaftaran."),
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(widths))
    table.style = "Table Grid"
    table.autofit = False

    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)

    header = table.rows[0].cells
    for idx, text in enumerate(rows[0]):
        set_cell_text(header[idx], text, bold=True, color="FFFFFF")
        set_cell_shading(header[idx], "1F4D78")
        header[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])

    doc.add_paragraph()


def add_note(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    run.font.size = Pt(10)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    p2.runs[0].font.size = Pt(9)
    doc.add_paragraph()


def add_figure(doc: Document, key: str) -> None:
    filename, caption = SCREENSHOTS[key]
    image_path = ASSET_DIR / filename
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.3))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = "Caption"
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def add_numbered_steps(doc: Document, steps: list[str]) -> None:
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(85, 85, 85)
    caption.paragraph_format.space_after = Pt(8)


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Manual Pengguna Teknis Yayasan Cendekia")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(85, 85, 85)


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Manual Pengguna Teknis\nAplikasi Website Yayasan Cendekia")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Panduan operasional untuk pengguna publik dan administrator konten")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(85, 85, 85)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Versi dokumen: 1.0 | Tanggal penyusunan: 13 Juli 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph()
    add_note(
        doc,
        "Ruang Lingkup",
        "Dokumen ini disusun berdasarkan aplikasi yang berjalan secara lokal di http://127.0.0.1:8080 pada saat pemeriksaan. Screenshot diambil dari halaman aktual aplikasi dan diberi penanda merah untuk membantu pengguna mengenali tombol, menu, atau area kerja utama.",
    )

    doc.add_heading("1. Gambaran Umum Aplikasi", level=1)
    doc.add_paragraph(
        "Aplikasi Website Yayasan Cendekia digunakan sebagai pusat informasi digital untuk calon orang tua siswa, sekaligus sebagai CMS sederhana bagi admin yayasan. Pengunjung dapat melihat profil yayasan, memilih unit sekolah, membaca artikel, membuka FAQ, menghubungi admin, dan mengisi formulir pendaftaran. Admin dapat memperbarui artikel, identitas yayasan, branding, data unit sekolah, serta FAQ melalui dashboard."
    )
    add_figure(doc, "home")

    add_table(
        doc,
        [
            ["Area", "Fungsi Utama", "Pengguna"],
            ["Halaman publik", "Menampilkan informasi yayasan, unit sekolah, pendaftaran, artikel, FAQ, kontak, dan peta lokasi.", "Orang tua/calon siswa"],
            ["Dashboard admin", "Mengelola konten website tanpa mengubah kode aplikasi.", "Admin yayasan/sekolah"],
            ["Penyimpanan konten", "Menyimpan konten CMS secara aman di MySQL/MariaDB.", "Tim teknis"],
        ],
        [1.45, 3.65, 1.35],
    )

    doc.add_heading("2. Prasyarat dan Hak Akses", level=1)
    doc.add_paragraph(
        "Pengguna publik tidak memerlukan akun untuk membuka website. Akses admin memerlukan password dashboard. Pada lingkungan lokal, aplikasi menggunakan password bawaan apabila variabel ADMIN_PASSWORD belum diatur. Pada production, password dan session secret sebaiknya dikonfigurasi melalui environment variable agar tidak menggunakan nilai bawaan."
    )
    add_figure(doc, "admin_login")
    add_bullets(
        doc,
        [
            "Alamat lokal halaman admin: /admin.php.",
            "Password lokal bawaan: cendekia-admin, selama ADMIN_PASSWORD belum diganti.",
            "Sesi admin dilindungi CSRF token dan cookie autentikasi bertanda tangan.",
            "File gambar yang diunggah dibatasi hingga 5 MB dengan format JPG, JPEG, PNG, WEBP, atau GIF.",
        ],
    )

    doc.add_heading("3. Navigasi Halaman Publik", level=1)
    doc.add_paragraph(
        "Header website berisi menu Yayasan, unit sekolah Cirebon, Cabang Losari, Artikel, FAQ, Kontak, serta tombol Hubungi Kami. Tombol utama pada beranda mengarahkan pengguna ke daftar unit sekolah atau halaman kontak yayasan."
    )
    add_numbered_steps(
        doc,
        [
            "Buka halaman Beranda.",
            "Gunakan menu navigasi untuk memilih halaman yang dibutuhkan.",
            "Pilih Lihat Unit Sekolah untuk langsung menuju daftar sekolah.",
            "Gunakan Hubungi Kami apabila ingin membuka percakapan WhatsApp yayasan.",
        ],
    )
    add_figure(doc, "home")

    doc.add_heading("4. Modul Unit Sekolah", level=1)
    doc.add_paragraph(
        "Beranda menampilkan unit sekolah berdasarkan lokasi layanan. Area Cendekia Cirebon memuat TK/PAUD, SD, SMP, dan SMK. Area Cabang Losari memuat TKIT Cendekia 2 Losari dan SD IT Cendekia 2 Losari. Setiap kartu menyediakan tombol Lihat Sekolah dan Daftar."
    )
    add_figure(doc, "unit_cards")
    add_numbered_steps(
        doc,
        [
            "Pilih unit sekolah yang sesuai dengan jenjang calon siswa.",
            "Klik Lihat Sekolah untuk membuka profil lengkap unit.",
            "Klik Daftar untuk menuju bagian formulir pendaftaran pada halaman unit.",
        ],
    )

    doc.add_heading("4.1 Profil Unit Sekolah", level=2)
    doc.add_paragraph(
        "Halaman unit sekolah menampilkan nama sekolah, deskripsi singkat, tombol pendaftaran, tombol WhatsApp, program unggulan, fasilitas, kegiatan siswa, formulir pendaftaran, dan peta lokasi. Konten ini bersumber dari data unit sekolah yang dapat diperbarui melalui dashboard admin."
    )
    add_figure(doc, "unit_hero")

    doc.add_heading("4.2 Formulir Pendaftaran Online", level=2)
    doc.add_paragraph(
        "Form pendaftaran berada pada bagian Pendaftaran di setiap halaman unit. Data yang diminta dapat berbeda sesuai jenjang, misalnya data calon siswa, nomor WhatsApp, asal sekolah, data orang tua, dan alamat. Untuk sebagian unit, form mengarah ke Google Form yang telah dikonfigurasi pada data sekolah."
    )
    add_numbered_steps(
        doc,
        [
            "Masuk ke halaman unit sekolah.",
            "Klik Daftar Sekarang atau gulir ke bagian Pendaftaran.",
            "Isi seluruh kolom wajib sesuai data calon siswa dan orang tua.",
            "Klik Kirim Pendaftaran.",
            "Hubungi admin sekolah melalui WhatsApp apabila perlu verifikasi atau informasi lanjutan.",
        ],
    )
    add_figure(doc, "registration")
    add_note(
        doc,
        "Catatan",
        "Beberapa unit dapat menampilkan pemberitahuan bahwa pendaftaran online belum dibuka. Jika pesan tersebut muncul, pengguna diarahkan untuk menghubungi admin sekolah melalui WhatsApp.",
    )

    doc.add_heading("5. Modul Artikel", level=1)
    doc.add_paragraph(
        "Halaman Artikel menampilkan berita, kegiatan, dan informasi sekolah yang berstatus Published. Artikel berisi kategori, judul, ringkasan, gambar utama, penulis, tanggal, dan detail bacaan. Artikel yang masih berstatus Draft tidak ditampilkan di halaman publik."
    )
    add_numbered_steps(
        doc,
        [
            "Buka menu Artikel.",
            "Pilih kartu artikel yang ingin dibaca.",
            "Baca isi artikel pada halaman detail.",
            "Gunakan navigasi header untuk kembali ke halaman lain.",
        ],
    )
    add_figure(doc, "articles")

    doc.add_heading("6. Modul FAQ Pendaftaran", level=1)
    doc.add_paragraph(
        "Halaman FAQ membantu calon orang tua mendapatkan jawaban cepat sebelum menghubungi admin. Pertanyaan ditampilkan dalam bentuk panel yang dapat dibuka dan ditutup."
    )
    add_numbered_steps(
        doc,
        [
            "Buka menu FAQ.",
            "Klik pertanyaan yang ingin dibaca.",
            "Baca jawaban yang muncul di bawah pertanyaan.",
            "Gunakan tombol Tanya Admin apabila jawaban belum mencukupi.",
        ],
    )
    add_figure(doc, "faq")

    doc.add_heading("7. Modul Kontak", level=1)
    doc.add_paragraph(
        "Halaman Kontak menampilkan alamat yayasan, email, tombol WhatsApp Yayasan, peta lokasi, serta daftar kontak WhatsApp masing-masing unit sekolah. Modul ini menjadi jalur utama untuk konsultasi pendaftaran dan informasi lokasi."
    )
    add_numbered_steps(
        doc,
        [
            "Buka menu Kontak.",
            "Pilih WhatsApp Yayasan untuk pertanyaan umum.",
            "Gunakan peta lokasi untuk melihat posisi yayasan.",
            "Pada bagian Admin Sekolah, pilih WhatsApp sesuai unit yang dituju.",
        ],
    )
    add_figure(doc, "contact")

    doc.add_heading("8. Masuk ke Dashboard Admin", level=1)
    doc.add_paragraph(
        "Dashboard admin digunakan untuk mengelola konten website. Setelah login, admin akan melihat tab Artikel, Branding & Tampilan, Identitas & Kontak, Unit Sekolah, FAQ, serta tombol Keluar."
    )
    add_numbered_steps(
        doc,
        [
            "Buka /admin.php.",
            "Masukkan password admin.",
            "Klik Masuk Dashboard.",
            "Pastikan tab pengelolaan muncul setelah login berhasil.",
            "Klik Keluar setelah selesai bekerja untuk menutup sesi admin.",
        ],
    )
    add_figure(doc, "admin_login")

    doc.add_heading("9. Admin - Modul Artikel", level=1)
    doc.add_paragraph(
        "Modul Artikel digunakan untuk menambah, mengubah, menerbitkan, menyimpan sebagai draft, atau menghapus artikel. Artikel Published muncul pada halaman Beranda dan halaman Artikel, sedangkan Draft tetap tersimpan di dashboard."
    )
    add_figure(doc, "admin_articles")
    add_numbered_steps(
        doc,
        [
            "Buka tab Artikel.",
            "Isi Judul, Kategori, Status, Penulis, Ringkasan, dan Isi artikel.",
            "Tambahkan gambar utama melalui upload lokal atau URL gambar internet.",
            "Pilih Published jika artikel siap tampil di website; pilih Draft jika belum selesai.",
            "Klik Simpan Artikel.",
            "Untuk mengubah artikel lama, klik Edit pada daftar artikel, lakukan perubahan, lalu simpan.",
            "Untuk menghapus artikel, klik Hapus dan konfirmasi tindakan.",
        ],
    )

    doc.add_heading("10. Admin - Branding & Tampilan", level=1)
    doc.add_paragraph(
        "Modul Branding & Tampilan mengatur identitas visual dan teks antarmuka website. Admin dapat memilih warna tema, tipe logo, teks logo, logo gambar, sub-teks header, teks tombol kontak, slogan footer, dan teks copyright."
    )
    add_figure(doc, "admin_branding")
    add_numbered_steps(
        doc,
        [
            "Buka tab Branding & Tampilan.",
            "Pilih Warna Tema Utama sesuai kebutuhan identitas visual.",
            "Pilih Tipe Logo: Teks Sederhana atau Gambar Custom.",
            "Isi Teks Logo jika menggunakan logo teks.",
            "Upload logo gambar jika menggunakan logo custom.",
            "Perbarui teks header, tombol kontak, slogan footer, dan copyright.",
            "Klik Simpan Perubahan.",
        ],
    )

    doc.add_heading("11. Admin - Identitas & Kontak Yayasan", level=1)
    doc.add_paragraph(
        "Modul ini mengatur informasi yayasan yang tampil pada beranda, footer, halaman kontak, tombol WhatsApp, dan peta lokasi. Data yang diisi harus konsisten dengan informasi resmi yayasan."
    )
    add_figure(doc, "admin_identity")
    add_numbered_steps(
        doc,
        [
            "Buka tab Identitas & Kontak.",
            "Isi Nama Yayasan, Tagline, dan Deskripsi.",
            "Masukkan Nomor Telepon/WhatsApp dengan awalan 62 tanpa spasi atau tanda plus.",
            "Isi Alamat Email Utama dan Alamat Lengkap Yayasan.",
            "Masukkan link Google Maps Embed pada kolom Maps.",
            "Tambahkan foto utama halaman yayasan melalui upload lokal atau URL gambar.",
            "Klik Simpan Identitas Yayasan.",
        ],
    )

    doc.add_heading("12. Admin - Unit Sekolah", level=1)
    doc.add_paragraph(
        "Tab Unit Sekolah menampilkan seluruh unit yang tersedia. Admin dapat memilih salah satu unit untuk mengubah nama, deskripsi, warna aksen, nomor WhatsApp, link pendaftaran, peta, gambar hero, program unggulan, fasilitas, dan kegiatan siswa."
    )
    add_figure(doc, "admin_school_list")
    add_numbered_steps(
        doc,
        [
            "Buka tab Unit Sekolah.",
            "Pilih kartu unit yang akan diperbarui.",
            "Klik Edit Konten Unit Ini.",
        ],
    )

    doc.add_heading("12.1 Edit Konten Unit Sekolah", level=2)
    doc.add_paragraph(
        "Pada halaman edit unit, admin mengelola konten detail yang akan tampil pada halaman publik sekolah tersebut. Program unggulan dapat ditambah atau dihapus, sedangkan fasilitas dan kegiatan siswa diisi satu item per baris."
    )
    add_figure(doc, "admin_school_edit")
    add_numbered_steps(
        doc,
        [
            "Perbarui Nama Lengkap Sekolah dan Deskripsi Singkat.",
            "Pilih Tema Warna Visual untuk aksen tombol dan label.",
            "Isi Nomor WhatsApp Admin Sekolah.",
            "Masukkan Link Pendaftaran dan Link Google Maps Embed.",
            "Upload atau masukkan URL gambar hero sekolah.",
            "Edit Program Unggulan; gunakan Tambah Program jika diperlukan.",
            "Isi Fasilitas dan Kegiatan Siswa, masing-masing satu item per baris.",
            "Klik Simpan Perubahan Unit.",
        ],
    )

    doc.add_heading("13. Admin - FAQ Pendaftaran", level=1)
    doc.add_paragraph(
        "Modul FAQ digunakan untuk menambah, mengedit, dan menghapus pertanyaan yang tampil pada halaman FAQ publik. FAQ sebaiknya berisi pertanyaan yang sering muncul dari calon orang tua siswa."
    )
    add_figure(doc, "admin_faq")
    add_numbered_steps(
        doc,
        [
            "Buka tab FAQ.",
            "Isi kolom Pertanyaan dan Jawaban.",
            "Klik Simpan FAQ.",
            "Untuk mengubah FAQ lama, klik Edit pada daftar FAQ.",
            "Untuk menghapus FAQ, klik Hapus dan konfirmasi tindakan.",
        ],
    )

    doc.add_heading("14. Catatan Teknis Operasional", level=1)
    add_table(
        doc,
        [
            ["Topik", "Keterangan"],
            ["Penyimpanan", "Saat development dan production, data CMS dibaca langsung dari MySQL/MariaDB."],
            ["Upload gambar", "Gambar tersimpan di tabel cms_uploads dan dilayani melalui endpoint uploads."],
            ["Environment", "Atur DB_HOST, DB_NAME, DB_USER, DB_PASSWORD, ADMIN_PASSWORD, dan SESSION_SECRET di server; tidak ada kredensial bawaan."],
            ["Batas upload", "Maksimal 5 MB untuk JPG, JPEG, PNG, WEBP, atau GIF."],
        ],
        [1.8, 4.65],
    )
    add_note(
        doc,
        "Peringatan Operasional",
        "Jangan menyimpan password atau kredensial database dalam source code. Gunakan user database khusus dengan hak akses minimum.",
    )

    doc.add_heading("15. Checklist Penggunaan Harian Admin", level=1)
    add_bullets(
        doc,
        [
            "Login ke dashboard hanya dari perangkat yang aman.",
            "Periksa kembali status artikel sebelum menyimpan sebagai Published.",
            "Gunakan gambar yang jelas, relevan, dan tidak melanggar hak penggunaan.",
            "Pastikan nomor WhatsApp menggunakan format internasional 62.",
            "Uji halaman publik setelah mengubah data penting seperti link pendaftaran atau peta.",
            "Logout setelah selesai melakukan perubahan.",
        ],
    )

    doc.add_heading("16. Pemecahan Masalah Singkat", level=1)
    add_table(
        doc,
        [
            ["Kondisi", "Kemungkinan Penyebab", "Tindakan"],
            ["Tidak bisa login admin", "Password salah atau environment ADMIN_PASSWORD berbeda.", "Periksa password yang berlaku di environment tersebut."],
            ["Artikel tidak tampil", "Status masih Draft.", "Ubah status menjadi Published dan simpan ulang."],
            ["Gambar gagal diupload", "Ukuran/format file tidak sesuai.", "Gunakan JPG, PNG, WEBP, atau GIF maksimal 5 MB."],
            ["Peta tidak tampil", "URL Maps Embed salah.", "Gunakan link embed Google Maps, bukan link berbagi biasa."],
            ["Perubahan tidak tersimpan", "Konfigurasi MySQL salah atau tabel belum diimpor.", "Periksa DB_* dan import database/schema.sql."],
        ],
        [1.9, 2.15, 2.4],
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
